from django.utils.translation import ugettext_lazy as _
from django.utils.timezone import localtime
from django.utils import formats
from docx.shared import Cm

from .reservation_base import ReservationDocxRenderer, ReservationReport
from .utils import iso_to_dt


class ReservationListDocxRenderer(ReservationDocxRenderer):
    def render_one(self, document, reservation, renderer_context):
        begin = localtime(iso_to_dt(reservation['begin']))
        end = localtime(iso_to_dt(reservation['end']))

        resource = self.resources[reservation['resource']]
        time_str = '%s–%s' % tuple([formats.date_format(x, 'G.i') for x in (begin, end)])

        table = self.day_table
        if table is None:
            table = self.day_table = document.add_table(rows=0, cols=0)
            table.autofit = False
            # Four columns, with the last one left empty for notes
            for width in (Cm(4), Cm(10), Cm(2), Cm(2)):
                table.add_column(width)

        row = table.add_row()
        cells = row.cells

        p = cells[0].paragraphs[0]
        run = p.add_run()
        run.bold = True
        run.text = time_str

        host_name = reservation.get('host_name') or reservation.get('reserver_name')
        run = cells[0].add_paragraph(host_name).add_run()
        # Add a paragraph as padding
        cells[0].add_paragraph()

        p = cells[1].paragraphs[0]
        run = p.add_run()
        run.bold = True
        run.text = '%s / %s' % (resource.name, resource.unit.name)
        subject = reservation.get('event_subject')
        if subject:
            cells[1].add_paragraph(subject)
        nr_participants = reservation.get('number_of_participants')

        cells[2].text = str(nr_participants) if nr_participants is not None else ''

    def render_all(self, data, document, renderer_context):
        last_day = None
        self.day_table = None

        for idx, rv in enumerate(data):
            begin = localtime(iso_to_dt(rv['begin']))
            day = begin.date()
            if day != last_day:
                document.add_heading(formats.date_format(begin, r'D j.n.Y'), 3)
                last_day = day
                self.day_table = None

            self.render_one(document, rv, renderer_context)


class ReservationListReport(ReservationReport):
    renderer_classes = (ReservationListDocxRenderer,)

    def get_queryset(self):
        queryset = super().get_queryset()
        queryset = queryset.order_by('begin', 'resource__unit__name', 'resource__name')
        return queryset

    def get_filename(self, request, data):
        return '%s.docx' % (_('reservation-list'),)
