import datetime
import io

from django.utils.translation import ugettext_lazy as _
from django.utils import formats
from django.utils.timezone import localtime
from rest_framework import exceptions, renderers, response, serializers, status, views
from rest_framework.settings import api_settings
from docx import Document
from docx.shared import Pt

from resources.models import Reservation, Unit


class DocxRenderer(renderers.BaseRenderer):
    media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    format = 'docx'
    charset = None
    render_style = 'binary'

    def render(self, data, media_type=None, renderer_context=None):
        unit = data['unit']
        day = data.get('day', datetime.date.today())
        include_resources_without_reservations = data['include_resources_without_reservations']
        document = Document()
        first_resource = True

        for resource in unit.resources.all():
            reservations = Reservation.objects.filter(resource=resource, begin__date=day).order_by('resource', 'begin')
            reservation_count = reservations.count()

            if reservation_count == 0 and not include_resources_without_reservations:
                continue

            # every resource on it's own page, a bit easier to add linebreak here than at the end
            if not first_resource:
                document.add_page_break()
            else:
                first_resource = False

            h = document.add_heading('%s\t%s' % (resource.name, formats.date_format(day)), 1)
            h.paragraph_format.space_after = Pt(48)

            if reservation_count == 0:
                p = document.add_paragraph('- %s - ' % _('No reservations'))
                p.paragraph_format.space_before = Pt(24)

            for reservation in reservations:

                # the time
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(24)
                p.add_run(formats.time_format(localtime(reservation.begin)) + ' - ' +
                          formats.time_format(localtime(reservation.end))).bold = True

                # collect attributes from the reservation, skip empty ones
                attrs = [(field, getattr(reservation, field)) for field in (
                    'event_subject',
                    'number_of_participants',
                    # TODO moar fields
                ) if getattr(reservation, field)]

                if not attrs:
                    # this should not normally happen as event_subject and number_of_participants
                    # should be required fields
                    p = document.add_paragraph('- %s - ' % _('No information available'))
                    p.paragraph_format.space_before = Pt(24)
                    continue

                table = document.add_table(rows=0, cols=2)

                # build the attribute table
                for attr in attrs:
                    row_cells = table.add_row().cells
                    row_cells[0].text = Reservation._meta.get_field(attr[0]).verbose_name + ':'
                    row_cells[1].text = str(attr[1])

        output = io.BytesIO()
        document.save(output)

        return output.getvalue()


class ReportParamSerializer(serializers.Serializer):
    day = serializers.DateField(required=False)
    unit = serializers.CharField()
    include_resources_without_reservations = serializers.BooleanField(required=False, default=False)

    def validate_unit(self, value):
        try:
            unit = Unit.objects.get(id=value)
        except Unit.DoesNotExist:
            raise exceptions.ValidationError(
                _('Invalid pk "{pk_value}" - object does not exist.').format(pk_value=value)
            )
        return unit


class UnitEventsDayReport(views.APIView):
    renderer_classes = (DocxRenderer,)

    def get(self, request, format=None):
        serializer = ReportParamSerializer(data=request.query_params)
        if not serializer.is_valid():
            return response.Response(
                data=serializer.errors,
                status=status.HTTP_400_BAD_REQUEST,
            )
        return response.Response(serializer.validated_data)

    def finalize_response(self, request, response, *args, **kwargs):
        response = super().finalize_response(request, response, *args, **kwargs)

        # use the first renderer from settings to display errors
        if response.status_code != 200:
            first_renderer = api_settings.DEFAULT_RENDERER_CLASSES[0]()
            response.accepted_renderer = first_renderer
            response.accepted_media_type = first_renderer.media_type

        return response
