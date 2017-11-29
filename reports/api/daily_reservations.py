import datetime
import io
import pytz
from docx.shared import Pt, Cm

from django.utils.translation import get_language, ugettext_lazy as _
from django.utils import formats
from django.utils.timezone import localtime
from django.conf import settings
from rest_framework import exceptions

from resources.models import Reservation, Resource, Unit
from resources.api.resource import ResourceSerializer
from .base import BaseReport, DocxRenderer
from .utils import iso_to_dt


FALLBACK_LANGUAGE = settings.LANGUAGES[0][0]


class DailyReservationsDocxRenderer(DocxRenderer):
    def render(self, data, media_type=None, renderer_context=None):
        renderer_context = renderer_context or {}
        day = renderer_context.get('day')

        include_resources_without_reservations = renderer_context.get('include_resources_without_reservations')
        document = self.create_document()

        first_resource = True
        atleast_one_reservation = False

        current_language = get_language()

        def get_translated_prop(obj, prop_name):
            prop = obj.get(prop_name, {})
            val = prop.get(current_language)
            if not val:
                return prop.get(FALLBACK_LANGUAGE)
            return val

        for resource in data:
            reservations = [rv for rv in resource['reservations'] if rv['state'] == Reservation.CONFIRMED]
            reservation_count = len(reservations)
            if reservation_count == 0 and not include_resources_without_reservations:
                continue

            atleast_one_reservation = True

            # every resource on it's own page, a bit easier to add linebreak here than at the end
            if not first_resource:
                document.add_page_break()
            else:
                first_resource = False

            document.add_heading(get_translated_prop(resource, 'name'), 1)
            document.add_heading(formats.date_format(day, format='D j.n.Y'), 2)

            if reservation_count == 0:
                run = document.add_paragraph().add_run(_('No reservations.'))
                run.font.size = Pt(20)

            for reservation in reservations:
                # the time
                begin = iso_to_dt(reservation['begin'])
                end = iso_to_dt(reservation['end'])
                range_str = formats.time_format(localtime(begin)) + '–' + formats.time_format(localtime(end))
                time_paragraph = document.add_heading(range_str, 3)
                time_paragraph.paragraph_format.space_before = Cm(1)

                # collect attributes from the reservation, skip empty ones
                attrs = [(field, reservation.get(field)) for field in (
                    'event_subject',
                    'reserver_name',
                    'host_name',
                    'number_of_participants',
                ) if reservation.get(field)]

                if not attrs:
                    document.add_paragraph(_('No information available'))
                    continue

                table = document.add_table(rows=0, cols=2)
                # build the attribute table
                for attr in attrs:
                    row_cells = table.add_row().cells
                    row_cells[0].text = Reservation._meta.get_field(attr[0]).verbose_name + ':'
                    row_cells[1].text = str(attr[1])

        if not atleast_one_reservation:
            document.add_heading(_('No reservations'), 1)

        output = io.BytesIO()
        document.save(output)

        return output.getvalue()


class DailyReservationsReport(BaseReport):
    serializer_class = ResourceSerializer
    renderer_classes = (DailyReservationsDocxRenderer,)

    def get_queryset(self):
        return Resource.objects.all().order_by('unit__name', 'name')

    def get_serializer_context(self):
        context = super().get_serializer_context()
        context['start'] = self.start
        context['end'] = self.end
        return context

    def filter_queryset(self, queryset):
        params = self.request.query_params
        unit = params.get('unit', '').strip()
        if unit:
            try:
                unit = Unit.objects.get(id=unit)
            except Unit.DoesNotExist:
                raise exceptions.NotFound(
                    _('Unit "{pk_value}" does not exist.').format(pk_value=unit)
                )
            queryset = queryset.filter(unit=unit)

        resources = params.get('resource', '').strip()
        if resources:
            resource_ids = [x.strip() for x in resources.split(',')]
            queryset = queryset.filter(id__in=resource_ids)

        if not unit and not resources:
            raise exceptions.ParseError(_('Either unit or a valid resource id is required.'))

        if not queryset:
            raise exceptions.NotFound(_('No resources found'))

        if unit:
            tz = unit.time_zone
        else:
            tz = queryset.first().unit.time_zone
        tz = pytz.timezone(tz)

        day = params.get('day', '').strip()
        if day:
            try:
                day = datetime.datetime.strptime(day, "%Y-%m-%d").date()
            except ValueError:
                raise exceptions.ParseError('day must be of ISO format (YYYY-MM-DD)')
        else:
            day = datetime.date.today()
        self.day = day
        start = tz.localize(datetime.datetime.combine(day, datetime.time(0, 0)))
        end = start + datetime.timedelta(days=1)
        self.start = start
        self.end = end

        return queryset

    def get_renderer_context(self):
        context = super().get_renderer_context()
        params = self.request.query_params
        val = params.get('include_resources_without_reservations', '').lower() in ['true', '1', 't', 'y', 'yes']
        context['include_resources_without_reservations'] = val
        if hasattr(self, 'day'):
            context['day'] = self.day
        return context

    def get_filename(self, request, validated_data):
        return '%s-%s.docx' % (_('day-report'), self.day.isoformat())
