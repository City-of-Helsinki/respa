from django.utils.translation import get_language, pgettext_lazy, ugettext_lazy as _
from django.utils.timezone import localtime
from docx.shared import Cm

from resources.models import Reservation
from resources.models.utils import format_dt_range

from .reservation_base import ReservationDocxRenderer, ReservationReport
from .utils import iso_to_dt


class ReservationDetailsDocxRenderer(ReservationDocxRenderer):
    def render_one(self, document, reservation, renderer_context):
        begin = localtime(iso_to_dt(reservation['begin']))
        end = localtime(iso_to_dt(reservation['end']))

        resource = self.resources[reservation['resource']]
        time_str = format_dt_range(get_language(), begin, end)

        document.add_heading(resource.name, 1)
        document.add_heading(resource.unit.name, 3)
        time_paragraph = document.add_heading(time_str, 2)
        time_paragraph.paragraph_format.space_before = Cm(1)

        # collect attributes from the reservation, skip empty ones
        attrs = [(field, reservation.get(field)) for field in (
            'event_subject',
            'reserver_name',
            'reserver_phone_number',
            'host_name',
            'event_description',
            'number_of_participants',
            'participants',
        ) if reservation.get(field)]

        if attrs:
            table = document.add_table(rows=0, cols=2)

            for attr in attrs:
                row_cells = table.add_row().cells
                row_cells[0].text = Reservation._meta.get_field(attr[0]).verbose_name + ':'
                row_cells[1].text = str(attr[1])
        else:
            document.add_paragraph(_('No information available'))

        # 'has_catering_order' is only set if the user has permission to view
        # the catering orders
        if reservation.get('has_catering_order'):
            catering_order = self.catering_orders[reservation['id']]
            document.add_heading(pgettext_lazy('report', 'Catering order'), 2)

            table = document.add_table(rows=0, cols=2)

            for order_line in catering_order.order_lines.all():
                row_cells = table.add_row().cells
                row_cells[0].text = order_line.product.name
                row_cells[1].text = '%s %s' % (str(order_line.quantity), _('pcs.'))

            if catering_order.message:
                row_cells = table.add_row().cells
                row_cells[0].text = _('Message')
                row_cells[1].text = catering_order.message

            row_cells = table.add_row().cells
            row_cells[0].text = _('Invoicing data')
            row_cells[1].text = catering_order.invoicing_data

    def render_all(self, data, document, renderer_context):
        for idx, rv in enumerate(data):
            if idx != 0:
                document.add_page_break()
            self.render_one(document, rv, renderer_context)


class ReservationDetailsReport(ReservationReport):
    renderer_classes = (ReservationDetailsDocxRenderer,)

    def get_filename(self, request, data):
        return '%s.docx' % (_('reservation-details'),)
